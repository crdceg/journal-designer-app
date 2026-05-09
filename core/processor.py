# =========================================================
# JOURNAL DESIGNER APP
# Processor Core
# =========================================================

import os
import sys
import json
import tempfile

from docx import Document
from docxcompose.composer import Composer
from docx.shared import Inches
from docx.oxml.ns import qn


# =========================================================
# BASE DIRECTORIES
# =========================================================

if getattr(sys, "frozen", False):

    # EXE temp extraction folder
    APP_DIR = sys._MEIPASS

    # Real folder beside EXE
    ROOT_DIR = os.path.dirname(
        sys.executable
    )

else:

    APP_DIR = os.path.dirname(
        os.path.dirname(
            os.path.abspath(__file__)
        )
    )

    ROOT_DIR = APP_DIR

# =========================================================
# CONFIG
# =========================================================

CONFIG_FILE = os.path.join(
    APP_DIR,
    "config",
    "config.json"
)

with open(CONFIG_FILE, "r", encoding="utf-8") as f:

    config = json.load(f)

JOURNALS = config.get("journals", {})

ISSUES = config.get("issues", {})

TEMPLATES = config.get("templates", {})

COVERS_FOLDER = config.get(
    "covers_folder",
    "covers"
)

OUTPUT_FOLDER = config.get(
    "output_folder",
    "output"
)

output_root = os.path.join(
    ROOT_DIR,
    OUTPUT_FOLDER
)

os.makedirs(
    output_root,
    exist_ok=True
)


# =========================================================
# ISSUE CODES
# =========================================================

ISSUE_CODES = {
    "1": {
        "code": "jan",
        "month_en": "January",
        "month_ar": "يناير",
        "issue_num": 1,
    },

    "2": {
        "code": "apr",
        "month_en": "April",
        "month_ar": "أبريل",
        "issue_num": 2,
    },

    "3": {
        "code": "jul",
        "month_en": "July",
        "month_ar": "يوليو",
        "issue_num": 3,
    },

    "4": {
        "code": "oct",
        "month_en": "October",
        "month_ar": "أكتوبر",
        "issue_num": 4,
    },
}


# =========================================================
# TEXT REPLACEMENT
# =========================================================

def _replace_text_in_run_safe(
    run,
    replacements
):

    text = run.text

    for key, val in replacements.items():

        if key in text:

            text = text.replace(
                key,
                val
            )

    run.text = text


def replace_in_paragraph_preserve_format(
    paragraph,
    replacements
):

    if not paragraph.runs:
        return

    for run in paragraph.runs:

        _replace_text_in_run_safe(
            run,
            replacements
        )

    full_text = "".join([
        r.text for r in paragraph.runs
    ])

    for key in replacements:

        if key in full_text:

            first_run = paragraph.runs[0]

            font = first_run.font

            new_text = full_text

            for k, v in replacements.items():

                new_text = new_text.replace(
                    k,
                    v
                )

            for r in paragraph.runs:

                r.text = ""

            new_run = paragraph.add_run(
                new_text
            )

            try:

                if font.name:
                    new_run.font.name = font.name

                if font.size:
                    new_run.font.size = font.size

                new_run.font.bold = font.bold

                new_run.font.italic = font.italic

                new_run.font.underline = font.underline

                if (
                    font.color
                    and getattr(
                        font.color,
                        "rgb",
                        None
                    )
                ):
                    new_run.font.color.rgb = (
                        font.color.rgb
                    )

            except Exception:
                pass

            break


def replace_in_cell(
    cell,
    replacements
):

    for p in cell.paragraphs:

        replace_in_paragraph_preserve_format(
            p,
            replacements
        )

    for tbl in getattr(cell, "tables", []):

        for r in tbl.rows:

            for c in r.cells:

                replace_in_cell(
                    c,
                    replacements
                )


def replace_in_element_xml(
    element,
    replacements
):
    """
    Replace placeholders
    inside textboxes / shapes.
    """

    try:

        for node in element.iter():

            if node.tag == qn("w:t"):

                if node.text:

                    text = node.text

                    for k, v in replacements.items():

                        if k in text:

                            text = text.replace(
                                k,
                                v
                            )

                    node.text = text

    except Exception:
        pass


def replace_placeholders_in_doc(
    doc,
    replacements
):

    # Paragraphs
    for p in doc.paragraphs:

        replace_in_paragraph_preserve_format(
            p,
            replacements
        )

    # Tables
    for table in doc.tables:

        for row in table.rows:

            for cell in row.cells:

                replace_in_cell(
                    cell,
                    replacements
                )

    # Headers / Footers
    for section in doc.sections:

        # Headers
        for attr in [
            "header",
            "first_page_header",
            "even_page_header"
        ]:

            hdr = getattr(
                section,
                attr,
                None
            )

            if hdr:

                for p in hdr.paragraphs:

                    replace_in_paragraph_preserve_format(
                        p,
                        replacements
                    )

                for tbl in hdr.tables:

                    for row in tbl.rows:

                        for cell in row.cells:

                            replace_in_cell(
                                cell,
                                replacements
                            )

                replace_in_element_xml(
                    hdr._element,
                    replacements
                )

        # Footers
        for attr in [
            "footer",
            "first_page_footer",
            "even_page_footer"
        ]:

            ftr = getattr(
                section,
                attr,
                None
            )

            if ftr:

                for p in ftr.paragraphs:

                    replace_in_paragraph_preserve_format(
                        p,
                        replacements
                    )

                for tbl in ftr.tables:

                    for row in tbl.rows:

                        for cell in row.cells:

                            replace_in_cell(
                                cell,
                                replacements
                            )

                replace_in_element_xml(
                    ftr._element,
                    replacements
                )


# =========================================================
# COVER HANDLING
# =========================================================

def find_cover_path(
    journal_key,
    issue_key,
    year
):

    folder = os.path.join(
        APP_DIR,
        COVERS_FOLDER,
        journal_key
    )

    code = ISSUE_CODES.get(
        str(issue_key),
        {}
    ).get("code")

    if not code:
        return None

    for ext in (
        ".png",
        ".jpg",
        ".jpeg"
    ):

        path = os.path.join(
            folder,
            f"{code}{year}{ext}"
        )

        if os.path.isfile(path):

            return path

    return None


# =========================================================
# HELPERS
# =========================================================

def compute_volume(
    year,
    first_year
):

    try:

        return (
            int(year)
            - int(first_year)
            + 1
        )

    except Exception:

        return 1


# =========================================================
# DOCUMENT MERGING
# =========================================================

def merge_with_composer(
    template_doc_path,
    research_path,
    output_path,
    cover_path=None
):

    base = Document(template_doc_path)

    # Replace cover image
    if (
        cover_path
        and os.path.isfile(cover_path)
    ):

        try:

            section = base.sections[0]

            page_width = (
                section.page_width.inches
            )

            left_margin = (
                section.left_margin.inches
            )

            right_margin = (
                section.right_margin.inches
            )

            usable_width = (
                page_width
                - (left_margin + right_margin)
            )

            if usable_width <= 0:

                usable_width = page_width

            first_shape_found = False

            for p in base.paragraphs:

                for run in p.runs:

                    if (
                        "graphic"
                        in run._element.xml
                        and not first_shape_found
                    ):

                        p._element.remove(
                            run._element
                        )

                        new_run = p.add_run()

                        new_run.add_picture(
                            cover_path,
                            width=Inches(
                                usable_width
                            )
                        )

                        first_shape_found = True

                        break

                if first_shape_found:
                    break

            # Fallback
            if not first_shape_found:

                first_para = (
                    base.paragraphs[0]
                    if base.paragraphs
                    else base.add_paragraph()
                )

                run = first_para.add_run()

                run.add_picture(
                    cover_path,
                    width=Inches(
                        usable_width
                    )
                )

        except Exception as e:

            print(
                "Cover replace failed:",
                e
            )

    composer = Composer(base)

    research = Document(research_path)

    composer.append(research)

    composer.save(output_path)


# =========================================================
# MAIN PROCESS
# =========================================================

def process_job(
    journal_key,
    lang,
    issue_key,
    year,
    research_file
):

    # Validate journal
    if journal_key not in JOURNALS:

        raise FileNotFoundError(
            "المجلة غير موجودة في config.json"
        )

    journal_cfg = JOURNALS[journal_key]

    # Validate template
    template_rel = TEMPLATES.get(lang)

    if not template_rel:

        raise FileNotFoundError(
            "لم يتم تعريف قالب لهذه اللغة"
        )

    template_path = os.path.join(
        APP_DIR,
        template_rel
    )

    if not os.path.isfile(template_path):

        raise FileNotFoundError(
            f"القالب غير موجود:\n{template_path}"
        )

    # Validate issue
    if str(issue_key) not in ISSUE_CODES:

        raise ValueError(
            "قيمة العدد غير صحيحة"
        )

    year_i = int(year)

    # Volume
    volume = compute_volume(
        year_i,
        journal_cfg.get(
            "first_volume_year",
            year_i
        )
    )

    issue_num = ISSUE_CODES[
        str(issue_key)
    ]["issue_num"]

    month_text = (
        ISSUE_CODES[str(issue_key)][
            "month_ar"
        ]
        if lang == "ar"
        else ISSUE_CODES[str(issue_key)][
            "month_en"
        ]
    )

    # Placeholder replacements
    replacements = {

        "{JOURNAL_NAME}":
            (
                journal_cfg.get("name_ar")
                if lang == "ar"
                else journal_cfg.get("name_en")
            ),

        "{VOLUME}":
            str(volume),

        "{ISSUE}":
            str(issue_num),

        "{MONTH}":
            month_text,

        "{YEAR}":
            str(year_i),

        "{ISSN}":
            journal_cfg.get(
                "issn",
                ""
            ),

        "{URL}":
            journal_cfg.get(
                "url",
                ""
            ),

        "{VOLUME_ISSUE_DATE}":
            (
                f"المجلد {volume} | "
                f"العدد {issue_num} | "
                f"{month_text} {year_i}"
            )
            if lang == "ar"
            else
            (
                f"Volume {volume} | "
                f"Issue {issue_num} | "
                f"{month_text}. {year_i}"
            )
    }

    # Open template
    doc = Document(template_path)

    # Replace placeholders
    replace_placeholders_in_doc(
        doc,
        replacements
    )

    # Temporary file
    with tempfile.NamedTemporaryFile(
        delete=False,
        suffix=".docx"
    ) as tmp:

        temp_template = tmp.name

    doc.save(temp_template)

    # Find cover
    cover = find_cover_path(
        journal_key,
        issue_key,
        str(year_i)
    )

    # Output file name
    research_basename = os.path.splitext(
        os.path.basename(research_file)
    )[0]

    safe_basename = (
        research_basename
        .replace("/", "_")
        .replace("\\", "_")
    )

    output_name = (
        f"{safe_basename} - مصمم الصفحة.docx"
    )

    output_root = os.path.join(
            ROOT_DIR,
            OUTPUT_FOLDER
        )

    os.makedirs(
            output_root,
            exist_ok=True
        )

    output_path = os.path.join(
            output_root,
            output_name
        )

            # Merge documents
    merge_with_composer(
                temp_template,
                research_file,
                output_path,
                cover
            )

    # Cleanup
    try:

        os.remove(temp_template)

    except Exception:
        pass

    return output_path